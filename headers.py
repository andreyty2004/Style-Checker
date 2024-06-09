from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH #чтобы смотреть выравнивания
import re
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import shutil
import os
from lxml import etree
from bs4 import BeautifulSoup

def headers(fpath = ""):
    doc = Document(fpath)
    output = "-- Проверка заголовков --\n"

    #нашли все разделы (заголовки первого уровня)
    for par in doc.paragraphs:
        if 'Heading 1' in par.style.name and len(par.runs) > 0:
            if (par.runs[0].contains_page_break) == False:
                output = output + "-> Раздел {0} должен начинаться с нового листа\n".format(par.text.upper())

    for i in range(0, len(doc.paragraphs)):
        if 'Heading 1' in doc.paragraphs[i].style.name and len(doc.paragraphs[i].runs) > 0:
            F = doc.paragraphs[i].style.font.bold #в случае если унаследовано от заголовка
            for j in range(0, len(doc.paragraphs[i].runs)):
                if doc.paragraphs[i].runs[j].bold != None:
                    F = doc.paragraphs[i].runs[j].bold
                if (F == True):
                    output = output + "-> Название раздела {0} не должно быть написано жирным шрифтом\n".format(doc.paragraphs[i].text.upper())
                    break

    for i in range(0, len(doc.paragraphs)):
        if 'Heading 1' in doc.paragraphs[i].style.name and len(doc.paragraphs[i].runs) > 0:
                for j in range(0, len(doc.paragraphs[i].runs)):
                    if doc.paragraphs[i].runs[j].font.italic == True:
                        output = output + "-> Название раздела {0} не должно быть написано курсивом\n".format(doc.paragraphs[i].text.upper())
                        break

    for i in range(0, len(doc.paragraphs)):
        if 'Heading 1' in doc.paragraphs[i].style.name and len(doc.paragraphs[i].runs) > 0:
                for j in range(0, len(doc.paragraphs[i].runs)):
                    if doc.paragraphs[i].runs[j].font.underline == True:
                        output = output + "-> Название раздела {0} не должно быть подчеркнуто\n".format(doc.paragraphs[i].text.upper())
                        break

    for i in range(0, len(doc.paragraphs)):
        if 'Heading 1' in doc.paragraphs[i].style.name and len(doc.paragraphs[i].runs) > 0: #второе условие довольно странное но вроде как проверка на пустоту
            size = doc.paragraphs[i].style.font.size.pt
            for j in range(0, len(doc.paragraphs[i].runs)):
                if doc.paragraphs[i].runs[j].font.size != None:
                    size = doc.paragraphs[i].runs[j].font.size.pt
                if size != 14.0:
                    output = output + "-> Название раздела {0} должно быть написано шрифтом 14 пунктов\n".format(doc.paragraphs[i].text.upper())
                    break

    for i in range(0, len(doc.paragraphs)):
        if 'Heading 1' in doc.paragraphs[i].style.name and len(doc.paragraphs[i].runs) > 0: #второе условие довольно странное но вроде как проверка на пустоту
            if doc.paragraphs[i].alignment != WD_ALIGN_PARAGRAPH.JUSTIFY:
                output = output + "-> Название раздела {0} должно быть выравнено по ширине\n".format(doc.paragraphs[i].text.upper())

    for i in range(0, len(doc.paragraphs)):
        if 'Heading 1' in doc.paragraphs[i].style.name and len(doc.paragraphs[i].runs) > 0: #второе условие довольно странное но вроде как проверка на пустоту
            soup = BeautifulSoup(doc.paragraphs[i]._p.xml, "xml")
            inds = soup.find_all("w:ind")
            if len(inds) != 0:
                if inds[0].get("w:left") != None:
                    indent = round((int(inds[0].get("w:left"))) / (20 * 28.346), 2)
                    if indent != 1.25:
                        output = output + "-> В названии раздела {0} должен быть задан абзацный отступ 1.25 см\n".format(doc.paragraphs[i].text.upper())
                else:
                    output = output + "-> В названии раздела {0} должен быть задан абзацный отступ 1.25 см\n".format(doc.paragraphs[i].text.upper())
            else:
                output = output + "-> В названии раздела {0} должен быть задан абзацный отступ 1.25 см\n".format(doc.paragraphs[i].text.upper())

    for i in range(0, len(doc.paragraphs)):
        if 'Heading 1' in doc.paragraphs[i].style.name and len(doc.paragraphs[i].runs) > 0: #второе условие довольно странное но вроде как проверка на пустоту
            string = doc.paragraphs[i].text
            if string[-1] == '.':
                output = output + "-> В конце названия раздела {0} точка не ставится\n".format(string.upper())
            if len(re.findall('(\d+.)', string)) > 0:
                substr = re.findall('(\d+.)', string)[0]
                if string.find(substr) == 0 and substr[-1] == '.':
                    output = output + "-> После номера раздела {0} точка не ставится\n".format(string.upper())
            

    #Попробовали посуетить с точкой между предложениями в заголовки

    for i in range(0, len(doc.paragraphs)):
        if 'Heading 1' in doc.paragraphs[i].style.name and len(doc.paragraphs[i].runs) > 0: #второе условие довольно странное но вроде как проверка на пустоту
            if (doc.paragraphs[i + 1].text == "" and doc.paragraphs[i + 2].text != "") == False:
                output = output + "-> Текст раздела должен отделяться от заголовка {0} пустой строкой\n".format(doc.paragraphs[i].text.upper())


    for i in range(0, len(doc.paragraphs)):
        if 'Heading 2' in doc.paragraphs[i].style.name and len(doc.paragraphs[i].runs) > 0:
            F = doc.paragraphs[i].style.font.bold
            for j in range(0, len(doc.paragraphs[i].runs)):
                if doc.paragraphs[i].runs[j].bold != None:
                    F = doc.paragraphs[i].runs[j].bold
                if F == True:
                    output = output + "-> Название подраздела {0} не должно быть написано жирным шрифтом\n".format(doc.paragraphs[i].text.upper())
                    break

    for i in range(0, len(doc.paragraphs)):
        if 'Heading 2' in doc.paragraphs[i].style.name and len(doc.paragraphs[i].runs) > 0:
            F = doc.paragraphs[i].style.font.italic
            for j in range(0, len(doc.paragraphs[i].runs)):
                if doc.paragraphs[i].runs[j].italic != None:
                    F = doc.paragraphs[i].runs[j].italic
                if F == True:
                    output = output + "-> Название подраздела {0} не должно быть написано курсивом\n".format(doc.paragraphs[i].text.upper())
                    break

    for i in range(0, len(doc.paragraphs)):
        if 'Heading 2' in doc.paragraphs[i].style.name and len(doc.paragraphs[i].runs) > 0:
                for j in range(0, len(doc.paragraphs[i].runs)):
                    if doc.paragraphs[i].runs[j].underline == True:
                        output = output + "-> Название подраздела {0} не должно быть подчеркнуто\n".format(doc.paragraphs[i].text.upper())
                        break

    for i in range(0, len(doc.paragraphs)):
        if 'Heading 2' in doc.paragraphs[i].style.name and len(doc.paragraphs[i].runs) > 0: #второе условие довольно странное но вроде как проверка на пустоту
            indent = round(doc.paragraphs[i].style.paragraph_format.first_line_indent.cm, 2)
            if doc.paragraphs[i].paragraph_format.first_line_indent != None:
                indent = round(doc.paragraphs[i].paragraph_format.first_line_indent.cm, 2)
            if indent != 1.25:
                output = output + "-> В названии подраздела {0} должен быть задан абзацный отступ 1.25 см\n".format(doc.paragraphs[i].text.upper())
            
    for i in range(0, len(doc.paragraphs)):
        if 'Heading 2' in doc.paragraphs[i].style.name and len(doc.paragraphs[i].runs) > 0: #второе условие довольно странное но вроде как проверка на пустоту
            size = doc.paragraphs[i].style.font.size.pt
            for j in range(0, len(doc.paragraphs[i].runs)):
                if doc.paragraphs[i].runs[j].font.size != None:
                    size = doc.paragraphs[i].runs[j].font.size.pt
                if size != 14.0:
                    output = output + "-> Название подраздела {0} должно быть написано шрифтом 14 пунктов\n".format(doc.paragraphs[i].text.upper())
                    break

    for i in range(0, len(doc.paragraphs)):
        if 'Heading 2' in doc.paragraphs[i].style.name and len(doc.paragraphs[i].runs) > 0: #второе условие довольно странное но вроде как проверка на пустоту
            if doc.paragraphs[i].alignment != WD_ALIGN_PARAGRAPH.JUSTIFY:
                output = output + "-> Название подраздела {0} должно быть выравнено по ширине\n".format(doc.paragraphs[i].text.upper())

    for i in range(0, len(doc.paragraphs)):
        if 'Heading 2' in doc.paragraphs[i].style.name and len(doc.paragraphs[i].runs) > 0: #второе условие довольно странное но вроде как проверка на пустоту
            if (doc.paragraphs[i + 1].text == "" and doc.paragraphs[i + 2].text != "") == False:
                output = output + "-> Текст подраздела должен отделяться от заголовка {0} пустой строкой\n".format(doc.paragraphs[i].text.upper())

    for i in range(0, len(doc.paragraphs)):
        if 'Heading 2' in doc.paragraphs[i].style.name and len(doc.paragraphs[i].runs) > 0: #второе условие довольно странное но вроде как проверка на пустоту
            string = doc.paragraphs[i].text
            if string[-1] == '.':
                output = output + "-> В конце названия подраздела {0} точка не ставится\n".format(string.upper())
            if len(re.findall('(\d+\.\d+.)', string)) > 0:
                substr = re.findall('(\d+\.\d+.)', string)[0]
                if string.find(substr) == 0 and substr[-1] == '.':
                    output = output + "-> После номера подраздела {0} точка не ставится\n".format(string.upper())


    for i in range(0, len(doc.paragraphs)):
        if 'Heading 1' in doc.paragraphs[i].style.name and len(doc.paragraphs[i].runs) > 0:
            head_str = doc.paragraphs[i].text
            c = i + 1
            number_of_subheadings = 0
            index = 0
            while ('Heading 1' in doc.paragraphs[c].style.name and len(doc.paragraphs[c].runs) > 0) == False:

                if 'Heading 2' in doc.paragraphs[c].style.name and len(doc.paragraphs[c].runs) > 0:
                    if (number_of_subheadings == 0):
                        index = c
                    number_of_subheadings += 1
                c += 1
                if c == len(doc.paragraphs):
                    break

            if (number_of_subheadings == 1):
                single = doc.paragraphs[index].text
                if len(re.findall('(\d+)', single)) != 0:
                    subsingle = re.findall('(\d+)', single)[0]
                    if single.find(subsingle) == 0:
                        output = output + "-> Единственный подраздел {0} в разделе {1} не нумеруется\n".format(single.upper(), head_str.upper())
                        continue


            
            if len(re.findall('(\d+)', head_str)) == 0:  
                continue
            head_substr = re.findall('(\d+)', head_str)[0]
            if head_str.find(head_substr) != 0:
                continue
            c = i + 1
            remembered = 0
            while ('Heading 1' in doc.paragraphs[c].style.name and len(doc.paragraphs[c].runs) > 0) == False:

                if 'Heading 2' in doc.paragraphs[c].style.name and len(doc.paragraphs[c].runs) > 0:
                    string = doc.paragraphs[c].text
                    if len(re.findall('(\d+\.\d+)', string)) > 0:
                        substr = re.findall('(\d+\.\d+)', string)[0]
                        if string.find(substr) == 0:
                            if substr.split('.')[0] != head_substr:
                                output = output + "-> Нумерация подраздела {0} должна проводиться в пределах основого раздела {1} (первая цифра - цифра раздела)\n".format(string.upper(), head_str.upper())
                            if int(substr.split('.')[1]) != remembered + 1:
                                output = output + "-> Нарушена последовательная нумерация подраздела {0} в разделе {1}\n".format(string.upper(), head_str.upper())
                            remembered = int(substr.split('.')[1])
                        else:
                            output = output + "-> Отсутствует нумерация подраздела {0} в разделе {1}\n".format(string.upper(), head_str.upper())
                    else:
                        output = output + "-> Отсутствует нумерация подраздела {0} в разделе {1}\n".format(string.upper(), head_str.upper())
                c += 1
                if c == len(doc.paragraphs):
                    break

    for i in range(0, len(doc.paragraphs)):
        if 'Heading 2' in doc.paragraphs[i].style.name and len(doc.paragraphs[i].runs) > 0:
            head_str = doc.paragraphs[i].text
            c = i + 1
            number_of_subheadings = 0
            index = 0
            while ('Heading 2' in doc.paragraphs[c].style.name and len(doc.paragraphs[c].runs) > 0) == False:

                if 'Heading 3' in doc.paragraphs[c].style.name and len(doc.paragraphs[c].runs) > 0:
                    if (number_of_subheadings == 0):
                        index = c
                    number_of_subheadings += 1
                c += 1
                if c == len(doc.paragraphs):
                    break

            if (number_of_subheadings == 1):
                single = doc.paragraphs[index].text
                if len(re.findall('(\d+)', single)) != 0:
                    subsingle = re.findall('(\d+)', single)[0]
                    if single.find(subsingle) == 0:
                        output = output + "-> Единственный пункт {0} в подразделе {1} не нумеруется\n".format(single.upper(), head_str.upper())
    
    if(output == "-- Проверка заголовков --\n"):
        output = output + "-> OK\n"
    output = output + '\n'
    return output