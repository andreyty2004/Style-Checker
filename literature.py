import re
from docx import Document
from lxml import etree
import zipfile
from docx.enum.text import WD_ALIGN_PARAGRAPH


def literature(fpath = ""):

    doc = Document(fpath)
    output = "-- Проверка списков источников --\n"

    string = re.findall('(xmlns:w=".*?")', doc._element.xml)[0]
    #? - ленивый захват (как можно меньшую строку)
    #. - любой символ
    #* - любое количество раз
    global w
    w = (re.findall('".*"', string)[0])[1:-1]
    #w - не должна зависеть от версии word

    global numbering_xml
    with zipfile.ZipFile(fpath, 'r') as zip_file:
        numbering_xml = zip_file.read('word/numbering.xml')



    index = 0
    for i in range(0, len(doc.paragraphs)):
        par = doc.paragraphs[i]
        if 'Heading 1' in par.style.name and len(par.runs) > 0:
            head_str = par.text
            if head_str.lower().count("список использованных источников") != 0:
                index = i
                break
    if index == 0:
        output = output + '-> Не найден заголовок "список использованных источников"\n'

    numbers = []
    for i in range(0, index):
        par = doc.paragraphs[i]
        text = par.text
        referencies = re.findall('(\[\d+.*?\])', text)
        for ref in referencies:
            numbers.append(re.findall('(\d+)', ref)[0])

    figures = [int(num) for num in numbers]
    figures1 = figures.copy()
    figures.sort()
    if figures1 != figures:
        output = output + '-> Список использованных источников следует формировать в порядке упоминания источников в тексте ВКРБ\n'

    cmp = ""
    last_number = 0
    alignment_checked = False

    for i in range(index + 1, len(doc.paragraphs)):
        par = doc.paragraphs[i]
        if 'Heading 1' in par.style.name and len(par.runs) > 0:
            break

        if par.text == "":
            continue

        if num(par) != None:
            # логика связанная с обработкой списков
            if par.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY and alignment_checked == False:
                output = output + f'-> текст списка использованных источников должен выравниваться по ширине\n' 
                alignment_checked = True

            if get_lvl_text(num(par), lvl(par)) != "%1." or get_num_fmt(num(par), lvl(par)) != 'decimal':
                output = output + "-> список использованных источников нумеруется арабскими цифрами без скобок, кавычек и других маркеров\n"
                break

        else:

            if par.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY and alignment_checked == False:
                output = output + f'-> текст списка использованных источников должен выравниваться по ширине\n'
                alignment_checked = True

            figures = re.findall('(\d+)', par.text)
            if len(figures) != 0:
                figure = figures[0]
                if par.text.find(figure) != 0:
                    output = output + "-> список использованных источников нумеруется арабскими цифрами без скобок, кавычек и других маркеров\n"
                    break
            else:
                output = output + "-> список использованных источников нумеруется арабскими цифрами без скобок, кавычек и других маркеров\n"
                break
            

            numbers = re.findall('(\d+[\.\t ])', par.text)
            if len(numbers) != 0:
                number = numbers[0]
                if par.text.find(number) != 0:
                    output = output + "-> список использованных источников нумеруется арабскими цифрами без скобок, кавычек и других маркеров\n"
                    break
            else:
                output = output + "-> список использованных источников нумеруется арабскими цифрами без скобок, кавычек и других маркеров\n"
                break

            save = re.findall('(\d+.*?[а-яА-Яa-zA-Z])', par.text)[0]
            length = len(re.findall('(\d+)', par.text)[0])
            save = save[length:-1]
            if cmp == "":
                cmp = save
            if cmp != save:
                output = output + "-> при нумерации списка использованных источников должен использоваться единый стиль нумерации\n"
                break


            number = int(re.findall('(\d+)', par.text)[0])
            if number == last_number + 1:
                last_number = number
            else:
                output = output + "-> Нарушен порядок нумерации при нумерации списка использованных источников\n"
                break
    if(output == "-- Проверка списков источников --\n"):
        output = output + "-> OK\n"
    output = output + '\n'
    return output


def get_num_fmt(numId, ilvl):
    tree = etree.fromstring(numbering_xml)

    abstractNumId = str(tree.xpath('w:num[@w:numId="{0}"]/w:abstractNumId/@w:val'.format(numId), namespaces = {'w': '{0}'.format(w)})[0])

    if len(tree.xpath('w:abstractNum[@w:abstractNumId="{0}"]/w:lvl[@w:ilvl="{1}"]/w:numFmt/@w:val'.format(abstractNumId, ilvl), namespaces = {'w': '{0}'.format(w)})) == 0:
        return None
    else:
        numFmt = str(tree.xpath('w:abstractNum[@w:abstractNumId="{0}"]/w:lvl[@w:ilvl="{1}"]/w:numFmt/@w:val'.format(abstractNumId, ilvl), namespaces = {'w': '{0}'.format(w)})[0])
        return numFmt

def get_lvl_text(numId, ilvl):
    tree = etree.fromstring(numbering_xml)
    abstractNumId = str(tree.xpath('w:num[@w:numId="{0}"]/w:abstractNumId/@w:val'.format(numId), namespaces = {'w': '{0}'.format(w)})[0])

    if len(tree.xpath('w:abstractNum[@w:abstractNumId="{0}"]/w:lvl[@w:ilvl="{1}"]/w:lvlText/@w:val'.format(abstractNumId, ilvl), namespaces = {'w': '{0}'.format(w)})) == 0:
        return None
    else:
        lvlText = str(tree.xpath('w:abstractNum[@w:abstractNumId="{0}"]/w:lvl[@w:ilvl="{1}"]/w:lvlText/@w:val'.format(abstractNumId, ilvl), namespaces = {'w': '{0}'.format(w)})[0])
        return lvlText

def num(paragraph):
    external_string = re.findall('(<w:numId.*?>)', paragraph.paragraph_format._element.xml)
    if len(external_string) != 0:
        external_string = re.findall('(<w:numId.*?>)', paragraph.paragraph_format._element.xml)[0]
        numId = re.findall('(\d+)', external_string)[0]
        return numId
    return None

def lvl(paragraph):
    external_string = re.findall('(<w:ilvl.*?>)', paragraph.paragraph_format._element.xml)[0]
    ilvl = re.findall('(\d+)', external_string)[0]
    return ilvl
        


        





    

























