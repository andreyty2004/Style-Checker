from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH #чтобы смотреть выравнивания
import re
from lxml import etree
import zipfile

letter_numbering = ['upperLetter', 'lowerLetter', 'russianUpper', 'russianLower']

def listings(fpath = ""):

    doc = Document(fpath)
    output = "-- Проверка списков --\n"

    string = re.findall('(xmlns:w=".*?")', doc._element.xml)[0]
    #? - ленивый захват (как можно меньшую строку)
    #. - любой символ
    #* - любое количество раз
    global w
    w = (re.findall('".*"', string)[0])[1:-1]
    #w - не должна зависеть от версии word

    global numbering_xml
    with zipfile.ZipFile(fpath, 'r') as zip_file:
        numbering_xml = zip_file.read('word/numbering.xml')

    numbered_paragraphs = []
    for i in range(0, len(doc.paragraphs)):
        if doc.paragraphs[i].style.name == 'List Paragraph':
            numbered_paragraphs.append(doc.paragraphs[i])

    single_numered_list = []
    if len(numbered_paragraphs) > 0:
        single_numered_list.append(numbered_paragraphs[0])
    depth = 1

    for i in range(0, len(numbered_paragraphs) - 1):
        if num(numbered_paragraphs[i]) == num(numbered_paragraphs[i + 1]):
            single_numered_list.append(numbered_paragraphs[i + 1])
            if lvl(numbered_paragraphs[i]) != lvl(numbered_paragraphs[i + 1]):
                depth += 1  
        else:
            proccesing(depth, single_numered_list)
            single_numered_list.clear()
            single_numered_list.append(numbered_paragraphs[i + 1])
            depth = 1

    output = output + proccesing(depth, single_numered_list)
    return output



def get_num_fmt(numId, ilvl):
    tree = etree.fromstring(numbering_xml)

    abstractNumId = str(tree.xpath('w:num[@w:numId="{0}"]/w:abstractNumId/@w:val'.format(numId), namespaces = {'w': '{0}'.format(w)})[0])

    if len(tree.xpath('w:abstractNum[@w:abstractNumId="{0}"]/w:lvl[@w:ilvl="{1}"]/w:numFmt/@w:val'.format(abstractNumId, ilvl), namespaces = {'w': '{0}'.format(w)})) == 0:
        return None
    else:
        numFmt = str(tree.xpath('w:abstractNum[@w:abstractNumId="{0}"]/w:lvl[@w:ilvl="{1}"]/w:numFmt/@w:val'.format(abstractNumId, ilvl), namespaces = {'w': '{0}'.format(w)})[0])
        return numFmt


def get_lvl_text(numId, ilvl):
    tree = etree.fromstring(numbering_xml)
    abstractNumId = str(tree.xpath('w:num[@w:numId="{0}"]/w:abstractNumId/@w:val'.format(numId), namespaces = {'w': '{0}'.format(w)})[0])

    if len(tree.xpath('w:abstractNum[@w:abstractNumId="{0}"]/w:lvl[@w:ilvl="{1}"]/w:lvlText/@w:val'.format(abstractNumId, ilvl), namespaces = {'w': '{0}'.format(w)})) == 0:
        return None
    else:
        lvlText = str(tree.xpath('w:abstractNum[@w:abstractNumId="{0}"]/w:lvl[@w:ilvl="{1}"]/w:lvlText/@w:val'.format(abstractNumId, ilvl), namespaces = {'w': '{0}'.format(w)})[0])
        return lvlText


def num(paragraph):
    external_string = re.findall('(<w:numId.*?>)', paragraph.paragraph_format._element.xml)[0]
    numId = re.findall('(\d+)', external_string)[0]
    return numId

def lvl(paragraph):
    external_string = re.findall('(<w:ilvl.*?>)', paragraph.paragraph_format._element.xml)[0]
    ilvl = re.findall('(\d+)', external_string)[0]
    return ilvl


def proccesing(depth, single_numered_list):
    ftext = ""
    # одноуровневый
    if depth == 1:
        for par in single_numered_list:
            if get_num_fmt(num(par), lvl(par)) not in ('decimal', 'bullet'):
                ftext = "-> При создании нумерованного одноуровневого списка {0} используются арабские цифры\n".format([par.text for par in single_numered_list])
                break

    # двухуровневый
    if depth == 2:
        for par in single_numered_list:
            if ((int(lvl(par)) == 0 and get_num_fmt(num(par), lvl(par)) != 'decimal') or
                (int(lvl(par)) == 1 and get_lvl_text(num(par), lvl(par)) != '-')):
                ftext = ftext + "-> При формировании двухуровневого списка {0} рекомендовано импользовать схему «номер – дефис»\n".format([par.text for par in single_numered_list])
                break

    # многоуровневый
    if depth > 2:
        for par in single_numered_list:
            if ((int(lvl(par)) == 0 and get_num_fmt(num(par), lvl(par)) != 'decimal') or
                (int(lvl(par)) == 1 and get_num_fmt(num(par), lvl(par)) not in letter_numbering) or
                (int(lvl(par)) == 2 and get_lvl_text(num(par), lvl(par)) != '-')):
                ftext = ftext + "-> Многоуровневые списки {0} рекомендуется создавать с соблюдением иерархии «номер – буква – дефис»\n".format([par.text for par in single_numered_list])
                break
    if(ftext == ""):
        ftext = "-> OK\n"
    ftext = ftext + "\n"
    return ftext
