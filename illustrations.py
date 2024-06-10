from docx import Document
import re
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup

def find_title(i, word):
    strs = [run.text.lower() for run in doc.paragraphs[i].runs]

    c = 1
    while c <= len(strs):
        for i in range(0, len(strs) - c + 1):
            string = ''
            for j in range(0, c):
                string += strs[i + j]
            if word in string:
                a = []
                for k in range(0, c):
                    a.append(i + k)        
                return a
        c += 1
    
    return []

def check_numbering_par(title_runs, number, c, name, is_appendix, page, par):
    ftext = ""
    title_text = ''.join([par.runs[i].text for i in range(title_runs[0], len(par.runs))])
    number_of_title = number + '.' + str(c)

    if number_of_title not in title_text:
        if is_appendix:
            ftext = f'-> Ошибка в подписи к рисунку "{name}" на странице {page}. Рисунки в приложениях нумеруются по схеме «номер приложения – точка – номер рисунка».\n'
        else:
            ftext = f'-> Ошибка в подписи к рисунку "{name}" на странице {page}. Рисунки в разделах нумеруются по схеме «номер раздела – точка – номер рисунка».\n'
    return ftext


def check_numbering_titles(stroka, number, c, name, is_appendix, page):
    ftext = ""
    soup = BeautifulSoup(stroka, "lxml")
    text = soup.find_all("w:t")
    title_text = ''.join([element.text for element in text])

    number_of_title = number + '.' + str(c)

    if number_of_title in title_text and title_text.find(number_of_title) == len(title_text) - len(number_of_title):
        pass
    else:
        if is_appendix:
            ftext = f'-> Ошибка в подписи к рисунку "{name}" на странице {page}. Рисунки в приложениях нумеруются по схеме «номер приложения – точка – номер рисунка».\n'
        else:
            ftext = ftext + f'-> Ошибка в подписи к рисунку "{name}" на странице {page}. Рисунки в разделах нумеруются по схеме «номер раздела – точка – номер рисунка».\n'
    return ftext


def page_increment(i, j):
    global isActive
    global page

    if j != len(doc.paragraphs[i].runs) - 1:
        if "w:br" in doc.paragraphs[i].runs[j]._r.xml and "w:lastRenderedPageBreak" in doc.paragraphs[i].runs[j + 1]._r.xml:
            page += 1
            isActive = False
            return 
        
    if i != len(doc.paragraphs) - 1 and len(doc.paragraphs[i + 1].runs) > 0 and j == len(doc.paragraphs[i].runs) - 1:
        if "w:br" in doc.paragraphs[i].runs[j]._r.xml and "w:lastRenderedPageBreak" in doc.paragraphs[i + 1].runs[0]._r.xml:
            page += 1
            isActive = False
            return

    if "w:br" in doc.paragraphs[i].runs[j]._r.xml:
        page += 1
        return

    if "w:lastRenderedPageBreak" in doc.paragraphs[i].runs[j]._r.xml and isActive:
        page += 1
        return

    if "w:lastRenderedPageBreak" in doc.paragraphs[i].runs[j]._r.xml:
        isActive = True  
        return


def get_picture_bottom_boarder_pt(run_xml):
    positionV = re.findall('(<wp:positionV(.|\n)*?wp:positionV>)', str(run_xml))
    if len(positionV) > 0:
        positionV = positionV[0]
    else:
        return None
    posOffset = re.findall('(<wp:posOffset.*?wp:posOffset>)', str(positionV))[0]
    top_boarder = re.findall('([-+]?\d+)', str(posOffset))[0]

    extent = re.findall('(<wp:extent(.|\n)*?>)', str(run_xml))[0]
    cy = re.findall('(cy=".*?")', str(extent))[0]
    height = re.findall('(\d+)', str(cy))[0]

    return (float(top_boarder) + float(height)) / 12700

def get_picture_name(run):
    string = str(run._r.xml)
    substr = re.findall('<wp:docPr.*?>', string)[0]
    substr1 = re.findall('descr=".*"', substr)[0]
    name = (re.findall('".*"', substr1)[0])[1:-1]
    return name

def get_title_top_margin(stroka):
    shape = re.findall('(<v:shape[^a-zA-Z].*?>)', stroka)[0]
    margin_top = re.findall('(margin-top:[-+]?\d+\.?\d+)', str(shape))[0]
    return float(re.findall('(\d+\.?\d+)', margin_top)[0])


def checking_str(stroka, name, bottom_boarder, page, in_one_par):
    ftext = ""

    if stroka.count('"center"') == 0:
        ftext = ftext + f'-> Подпись к рисунку "{name}" на странице {page} должна размещаться по центру\n'
    
    if stroka.count('<w:sz w:val="28"/>') == 0:
        ftext = ftext + f'-> Размер шрифра подписи к рисунку "{name}" на странице {page} должен совпадать с размером шрифта в основном тексте\n'

    if 'w:b' in stroka:
        if 'w:b w:val="0"' not in stroka:
            ftext = ftext + f'-> Шрифт подписи к рисунку "{name}" на странице {page} не должен быть написан жирным шрифтом\n'
        
    runs = BeautifulSoup(stroka, "lxml").find_all("w:r") # прогоны в подписи w:pict
    for tag in runs:
        if len(tag.find_all("w:i")) != 0:
            ftext = ftext + f'-> Шрифт подписи к рисунку "{name}" на странице {page} не должен быть написан курсивом\n'
            break
        

    for tag in runs:
        if len(tag.find_all("w:u")) != 0:
            ftext = ftext + f'-> Шрифт подписи к рисунку "{name}" на странице {page} не должен быть подчеркнут\n'
            break


    if 'w:hAnsi="Times New Roman"' not in stroka:
        ftext = ftext + f'-> Тип шрифра подписи к рисунку "{name}" на странице {page} должен совпадать с типом шрифта в основном тексте\n'

    if in_one_par and bottom_boarder != None:
        if abs(get_title_top_margin(stroka) - bottom_boarder) > 2.8:
            ftext = ftext + f'-> Подпись должна размещаться сразу под рисунком "{name}" на странице {page}\n'

    if stroka.count('stroked="f"') == 0:
        ftext = ftext + f'-> Подпись к рисунку {name} на странице {page} не должна выделяться рамкой\n'
    return ftext


def markup_options(run, name, page, string_list): #r - run     page - number of page

    sec = doc.sections[0]
    left_margin = sec.left_margin.pt
    right_margin = sec.right_margin.pt
    page_width = sec.page_width.pt

    if len(re.findall('(<wp:positionH(.|\n)*?wp:positionH>)', str(run._r.xml))) == 0:
        left_boarder = 0
    else:
        positionH = re.findall('(<wp:positionH(.|\n)*?wp:positionH>)', str(run._r.xml))[0]
        posOffset = re.findall('(<wp:posOffset.*?wp:posOffset>)', str(positionH))[0]
        left_boarder = re.findall('([-+]?\d+)', str(posOffset))[0]

    extent = re.findall('(<wp:extent(.|\n)*?>)', str(run._r.xml))[0]
    cx = re.findall('(cx=".*?")', str(extent))[0]
    picture_width = re.findall('(\d+)', str(cx))[0]

    cy = re.findall('(cy=".*?")', str(extent))[0]
    picture_heigth = re.findall('(\d+)', str(cy))[0]

    picture_area = (float(picture_width) / 12700) * (float(picture_heigth) / 12700)
    
    # 2.8 pt ~ 1 mm    достаточно близко
    if (('wp:wrapTopAndBottom' in str(run._r.xml)) or 
        ('wrapText="bothSides"' in str(run._r.xml) and abs(float(left_boarder) / 12700) < 2.8) and
        (abs((float(picture_width) / 12700 + float(left_boarder) / 12700 + left_margin) - (page_width - right_margin)) < 2.8)) == False:
        string_list[0] = f'-> При подготовке иллюстрации "{name}" на странице {page} в редакторе MSWord следует использовать опции меню «Параметры разметки – Обтекание текстом – Сверху и снизу» или «Параметры разметки – Обтекание текстом – Вокруг рамки» при условии, что ширина рамки совпадает с шириной текста\n'        
    
    return picture_area

def empty_string_separation(i, j, page): #i - index of drawing, j - index of title
    ftext = ""
    drawing_par = doc.paragraphs[i]
    if 'wp:wrapTopAndBottom' in str(drawing_par._p.xml) or 'wrapText="bothSides"' in str(drawing_par._p.xml):
        if len(doc.paragraphs[min(i, j) - 1].runs) != 0 and len(doc.paragraphs[max(i, j) + 1].runs) != 0:
            ftext = f'-> При размещении иллюстрации "{name}" на странице {page} в тексте следует отделять рисунок от текста пустой строкой и сверху, и снизу\n'
    return ftext            


def process_picture_area(i, j, picture_area, page, string_list): #i - drawing      j - title
    sec = doc.sections[0]
    page_width = sec.page_width.pt
    page_heigth = sec.page_height.pt
    page_area = page_width * page_heigth

    if (picture_area < 0.5 * page_area) and doc.paragraphs[i].runs[0].contains_page_break:
        for k in range(1, 4):
            if len(doc.paragraphs[max(i, j) + k].runs) > 0:
                if (doc.paragraphs[max(i, j) + k].runs[0].contains_page_break):
                    string_list[0] = f'-> Рисунок "{name}" на странице {page} слишком маленький, чтобы размещать его на отдельной странице\n'
                    return False

                if len(doc.paragraphs[max(i, j) + k].runs) > 1:
                    if ('w:br' in doc.paragraphs[max(i, j) + k].runs[0]._r.xml and doc.paragraphs[max(i, j) + k].runs[1].contains_page_break):
                        string_list[0] = string_list[0] + f'-> Рисунок "{name}" на странице {page} слишком маленький, чтобы размещать его на отдельной странице\n'
                        return False

                if len(doc.paragraphs[max(i, j) + k].runs) == 1 and len(doc.paragraphs[max(i, j) + k + 1].runs) > 0:
                    if ('w:br' in doc.paragraphs[max(i, j) + k].runs[0]._r.xml and doc.paragraphs[max(i, j) + k + 1].runs[0].contains_page_break):
                        string_list[0] = string_list[0] + f'-> Рисунок "{name}" на странице {page} слишком маленький, чтобы размещать его на отдельной странице\n'
                        return False
                return True

    if doc.paragraphs[i].runs[0].contains_page_break:
        return False

    return True

        
def checking_par(par, name, title_runs, page):
    ftext = ""

    if par.alignment != WD_ALIGN_PARAGRAPH.CENTER:
        ftext = f'-> Подпись к рисунку "{name}" на странице {page} должна размещаться по центру\n'

    size = par.style.font.size
    for j in title_runs:
        if par.runs[j].font.size != None:
            size = par.runs[j].font.size
        if size != None and size.pt != 14.0:
            ftext = ftext + f'-> Размер шрифра подписи к рисунку "{name}" на странице {page} должен совпадать с размером шрифта в основном тексте\n'
            break
    if size == None:
        ftext = ftext + f'-> Размер шрифра подписи к рисунку "{name}" на странице {page} должен совпадать с размером шрифта в основном тексте\n'

    F = par.style.font.bold #в случае если унаследовано от заголовка
    for j in title_runs:
        if par.runs[j].bold != None:
            F = par.runs[j].bold
        if (F == True):
            ftext = ftext + f'-> Шрифт подписи к рисунку "{name}" на странице {page} не должен быть написан жирным шрифтом\n'
            break

    for j in title_runs:
        if par.runs[j].font.italic == True:
            ftext = ftext + f'-> Шрифт подписи к рисунку "{name}" на странице {page} не должен быть написан курсивом\n'
            break


    for j in title_runs:
        if par.runs[j].font.name != 'Times New Roman':
            ftext = ftext + f'-> Тип шрифра подписи к рисунку "{name}" на странице {page} должен совпадать с типом шрифта в основном тексте\n'
            break

    for j in title_runs:
        if par.runs[j].font.underline == True:
            ftext = ftext + f'-> Шрифт подписи к рисунку "{name}" на странице {page} не должен быть подчеркнут\n'
            break

    return ftext
    
def check_numbering_above(i, name, number, c, page):
    ftext = ""
    par = doc.paragraphs[i]

    mention_runs = find_title(i, 'рисун')

    mention_text = ''.join([par.runs[i].text for i in range(mention_runs[0], len(par.runs))])
    number_of_title = number + '.' + str(c)

    if number_of_title not in mention_text:
        ftext = ftext + f'-> Ошибка в упоминании рисунка "{name}" на странице {page} в тексте перед рисунком. Рисунки в разделах нумеруются по схеме «номер раздела – точка – номер рисунка»\n'
    return ftext

def mentioned_above(i, name, number, c, page): # i - index of drawing
    ftext = ""
    par = doc.paragraphs[i]

    if 'Рис.' in par.text or 'рис.' in par.text:
        ftext = f'-> Ссылка на рисунок "{name}" на странице {page} должна быть дана без сокращений\n'
        return

    if 'рисун' in par.text.lower():
        ftext = ftext + check_numbering_above(i, name, number, c, page)

        if 'w:instrText' in par._p.xml:
            return
        else:
            ftext = ftext + f'-> Ссылка на рисунок "{name}" на странице {page} в тексте должна являться перекрестной ссылкой (при нажатии на нее переносит на рисунок)\n'
            return 
    
    
    for k in range(1, 4):
        par_k = doc.paragraphs[i - k]
        if len(par_k.runs) != 0:
            if (len(par_k.runs) == 1 and 'w:br' in par_k.runs[0]._r.xml):
                continue


            c1 = par_k.text.count('Рис.')
            c2 = par_k.text.count('рис.')
            if c1 > 0 or c2 > 0:
                ftext = ftext = f'-> Ссылка на рисунок "{name}" на странице {page} должна быть дана без сокращений\n'
                break

            c1 = par_k.text.count('Рисун')
            c2 = par_k.text.count('рисун')
            if c1 + c2 == 0:
                ftext = ftext + f'-> Не найдено упоминание рисунка "{name}" на странице {page} в тексте перед рисунком\n'
                break
            
            if c1 + c2 != 0:
                ftext = ftext + check_numbering_above(i - k, name, number, c, page)
                if 'w:instrText' not in par_k._p.xml:
                    ftext = ftext + f'-> Ссылка на рисунок "{name}" на странице {page} в тексте должна являться перекрестной ссылкой (при нажатии на нее переносит на рисунок)\n'

            break
    return ftext

def illustrations(fpath = ""):

    global page
    global name
    global doc
    global isActive
    doc = Document(fpath)
    output = "-- Проверка иллюстраций --\n"

    for i in range(0, len(doc.paragraphs)):
        count1 = doc.paragraphs[i].text.count('фиг.')
        count2 = doc.paragraphs[i].text.count('Фиг.')
        if count1 + count2 > 0:
            output = output + '-> В русскоязычной литературе не принято использовать обозначения «фигура 1», «Фиг.1».\n'
            break

    is_appendix = False #приложение
    number_defined = False #нумерация заголовка определена

    page = 1
    isActive = True # можно ли увеличивать page

    for i in range(0, len(doc.paragraphs)):

        par = doc.paragraphs[i]

        if 'Heading 1' in par.style.name and len(par.runs) > 0:
            number_defined = False
            head_str = par.text

            if 'приложение' in head_str.lower():
                is_appendix = True
                start = head_str.lower().find('приложение') + len('приложение')
                stroka = head_str[start:] #строка без 'приложение'
                numbers = re.findall('([А-Я])', stroka)
                if len(numbers) > 0:
                    number = numbers[0]
                    number_defined = True
                    c = 0
            else:
                numbers = re.findall('(\d+)', head_str)
                if len(numbers) > 0:
                    number = numbers[0]

                    if head_str.find(number) == 0:
                        number_defined = True
                        c = 0



        for j in range(0, len(par.runs)):
            run = par.runs[j]

            page_increment(i, j)

            if 'w:drawing' in str(run._r.xml): 
                if number_defined:
                    c += 1 #какой по счету рисунок в этом разделе     

                name = get_picture_name(run)
                temp_string_list = [""]
                picture_area = markup_options(run, name, page, temp_string_list)
                if(len(temp_string_list) != 0): 
                    output = output + temp_string_list[0]

                if is_appendix == False:
                    output = output + mentioned_above(i, name, number, c, page)

            
                if 'a:noFill' not in str(run._r.xml):
                    output = output + f'-> Рисунок "{name}" на странице {page} не должен выделяться рамкой\n'

                search = re.findall('(w:pict(.|\n)*?w:pict)', str(par._p.xml))
                length = len(search)

                #нумерован надписью
                if length > 0:
                    stroka = str(search[-1])
                    for k in reversed(range(0, length)):
                        if 'w:instr' in str(search[k]):
                            stroka = str(search[k])
                            break

                    bottom_boarder = get_picture_bottom_boarder_pt(run._r.xml)
                    my_string_list = [""]
                    output = output + checking_str(stroka, name, bottom_boarder, page, True)
                    if process_picture_area(i, i, picture_area, page, my_string_list):
                        empty_string_separation(i, i, page)
                        if(len(my_string_list) != 0):
                            output = output + my_string_list[0]
                    if number_defined:
                        #          номер_раздела.номер_рисунка
                        output = output + check_numbering_titles(stroka, number, c, name, is_appendix, page)
                    continue

                search = re.findall('(w:pict(.|\n)*?w:pict)', str(doc.paragraphs[i - 1]._p.xml))
                length = len(search)

                if length > 0:
                    stroka = str(search[-1])
                    for k in reversed(range(0, length)):
                        if 'w:instr' in str(search[k]):
                            stroka = str(search[k])
                            break

                    bottom_boarder = get_picture_bottom_boarder_pt(run._r.xml)
                    my_string_list = [""]
                    output = output + checking_str(stroka, name, bottom_boarder, page, False)
                    if process_picture_area(i, i - 1, picture_area, page, my_string_list):
                        empty_string_separation(i, i - 1, page)
                        output = output + my_string_list[0]
                    if number_defined:
                        output = output + check_numbering_titles(stroka, number, c, name, is_appendix, page)
                    continue


                search = re.findall('(w:pict(.|\n)*?w:pict)', str(doc.paragraphs[i + 1]._p.xml))
                length = len(search)
                if length > 0:
                    stroka = str(search[-1])
                    for k in reversed(range(0, length)):
                        if 'w:instr' in str(search[k]):
                            stroka = str(search[k])
                            break

                    bottom_boarder = get_picture_bottom_boarder_pt(run._r.xml)
                    my_string_list = [""]
                    output = output + checking_str(stroka, name, bottom_boarder, page, False)
                    if process_picture_area(i, i + 1, picture_area, page, my_string_list):
                        empty_string_separation(i, i + 1, page)
                        output = output + my_string_list[0]
                    if number_defined:
                        output = output + check_numbering_titles(stroka, number, c, name, is_appendix, page)
                    continue

                if doc.paragraphs[i + 1].text != '':
                    if len(title_runs := find_title(i + 1, 'рисунок')) != 0:
                        output = output + checking_par(doc.paragraphs[i + 1], name, title_runs, page)
                        my_string_list = [""]
                        if process_picture_area(i, i + 1, picture_area, page, my_string_list):
                            empty_string_separation(i, i + 1, page)
                            output = output + my_string_list[0]
                        if number_defined:
                            output = output + check_numbering_par(title_runs, number, c, name, is_appendix, page, doc.paragraphs[i + 1])
                        continue


                if doc.paragraphs[i + 2].text != '':
                    if len(title_runs := find_title(i + 2, 'рисунок')) != 0:
                        output = output + checking_par(doc.paragraphs[i + 2], name, title_runs, page)
                        my_string_list = [""]
                        if process_picture_area(i, i + 2, picture_area, page, my_string_list):
                            empty_string_separation(i, i + 2, page)
                            output = output + my_string_list[0]
                        if number_defined:
                            output = output + check_numbering_par(title_runs, number, c, name, is_appendix, page, doc.paragraphs[i + 2])
                        continue

                if doc.paragraphs[i].text != '':
                    if len(title_runs := find_title(i, 'рисунок')) != 0:
                        output = output + checking_par(par, name, title_runs, page)
                        my_string_list = [""]
                        if process_picture_area(i, i, picture_area, page, my_string_list):
                            empty_string_separation(i, i, page)
                            output = output + my_string_list[0]
                        if number_defined:
                            output = output + check_numbering_par(title_runs, number, c, name, is_appendix, page, doc.paragraphs[i])
                        continue

                output = output + f'-> Подпись к иллюстрации {name} на странице {page} не найдена\n'
    if(output == "-- Проверка иллюстраций --\n"):
        output = output + "-> OK\n" 
    output = output + '\n'
    return output

