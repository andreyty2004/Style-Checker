import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def applications(fpath = ""):

    doc = Document(fpath)
    output = "-- Проверка приложений --\n"

    index = -1
    for i in range(0, len(doc.paragraphs)):
        par = doc.paragraphs[i]
        if 'Heading 1' in par.style.name and len(par.runs) > 0:
            head_str = par.text
            if 'приложение' in head_str.lower():
                index = i
                break


    letters = []
    for i in range(0, index):
        par = doc.paragraphs[i]
        text = par.text

        apps = re.findall('([пП]риложени[^.]*? [А-ЯЁ][ )])', text)
        for t in range(0, len(apps)):
            string = apps[t]
            letters.append(string[-2])

    letters1 = letters.copy()
    letters.sort()

    if letters1 != letters:
        output = output + '-> Приложения должны нумероваться в порядке ссылок на них в тексте ВКРБ\n'

    wrong_letters = ['Ё', 'З', 'Й', 'О', 'Ч', 'Ь', 'Ы', 'Ъ']

    intersection = set(letters).intersection(wrong_letters)

    if len(intersection) != 0:
       output = output + '-> Приложения обозначают заглавными буквами русского алфавита, начиная с А, за исключением букв Ё, З, Й, О, Ч, Ь, Ы, Ъ\n'


    f = False
    for i in range(0, index):
        par = doc.paragraphs[i]
        if 'Heading 1' in par.style.name and len(par.runs) > 0:
            head_str = par.text
            if head_str.lower().count("список использованных источников") != 0:
                f = True
                break

    if f == False:
        output = output + '-> Приложения размещаются после списка использованных источников\n'

    if index != -1:
        for i in range(index, len(doc.paragraphs)):
            par = doc.paragraphs[i]
            if 'Heading 1' in par.style.name and len(par.runs) > 0:
                head_str = par.text
                if 'приложение' in head_str.lower():
                    if par.runs[0].contains_page_break == False:
                        output = output + f"-> {head_str} должно начинаться с новой страницы\n"

                    if par.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                        output = output + f'-> {head_str} должна размещаться по центру\n'

                    size = par.style.font.size
                    for j in range(0, len(par.runs)):
                        if par.runs[j].font.size != None:
                            size = par.runs[j].font.size
                        if size != None and size.pt != 14.0:
                            output = output + f'-> Размер шрифра {head_str} должен совпадать с размером шрифта в основном тексте\n'
                            break
                    if size == None:
                        output = output + f'-> Размер шрифра {head_str} должен совпадать с размером шрифта в основном тексте\n'
                    
                    F = par.style.font.bold #в случае если унаследовано от заголовка
                    for j in range(0, len(par.runs)):
                        if par.runs[j].bold != None:
                            F = par.runs[j].bold
                        if (F == True):
                            output = output + f'-> Шрифт {head_str} не должен быть написан жирным шрифтом\n'
                            break

                    for j in range(0, len(par.runs)):
                        if par.runs[j].font.italic == True:
                            output = output + f'-> Шрифт {head_str} не должен быть написан курсивом\n'
                            break


                    if 'w:hAnsi' in str(par._p.xml):
                        if 'w:hAnsi="Times New Roman"' not in str(par._p.xml):
                            output = output + f'-> Тип шрифра {head_str} должен совпадать с типом шрифта в основном тексте\n'
                    elif 'w:hAnsiTheme' in str(par._p.xml):
                        if 'w:hAnsiTheme="Times New Roman"' not in str(par._p.xml):
                            output = output + f'-> Тип шрифра {head_str} должен совпадать с типом шрифта в основном тексте\n'
                    elif 'w:cs' in str(par._p.xml):
                        if 'w:cs="Times New Roman"' not in str(par._p.xml):
                            output = output + f'-> Тип шрифра {head_str} должен совпадать с типом шрифта в основном тексте\n'



                    for j in range(0, len(par.runs)):
                        if par.runs[j].font.underline == True:
                            output = output + f'-> Шрифт {head_str} не должен быть подчеркнут\n'
                            break

                    if (head_str.lower().find("приложение") == 0 and "w:ind" not in str(par._p.xml)) == False:
                        output = output + f'-> Название {head_str} должно распологаться без абзацного отступа\n'


    alignment_checked = False
    size_checked = False
    bold_checked = False
    italic_checked = False
    font_checked = False
    underline_checked = False

    is_checking = True

    if index != -1:
        for i in range(index, len(doc.paragraphs)):
            par = doc.paragraphs[i]
            if 'Heading 1' in par.style.name and len(par.runs) > 0:
                head_str = par.text
                if 'приложение' in head_str.lower():
                    is_checking = True
                    alignment_checked = False
                    size_checked = False
                    bold_checked = False
                    italic_checked = False
                    font_checked = False
                    underline_checked = False
                else:
                    is_checking = False

                continue

            if is_checking:
                if par.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY and alignment_checked == False:
                    output = output + f'-> текст {head_str} должен размещаться по ширине\n' 
                    alignment_checked = True

                size = par.style.font.size
                for j in range(0, len(par.runs)):
                    if par.runs[j].font.size != None:
                        size = par.runs[j].font.size
                    if size != None and size.pt != 14.0 and size_checked == False:
                        output = output + f'-> Размер шрифра в тексте {head_str} должен совпадать с размером шрифта в основном тексте\n'
                        size_checked = True
                        break
                if size == None and size_checked == False:
                    output = output + f'-> Размер шрифра в тексте {head_str} должен совпадать с размером шрифта в основном тексте\n'
                    size_checked = True

                F = par.style.font.bold #в случае если унаследовано от заголовка
                for j in range(0, len(par.runs)):
                    if par.runs[j].bold != None:
                        F = par.runs[j].bold
                    if F == True and bold_checked == False:
                        output = output + f'-> Шрифт в тексте {head_str} не должен быть написан жирным шрифтом\n'
                        bold_checked = True

                for j in range(0, len(par.runs)):
                    if par.runs[j].font.italic == True and italic_checked == False:
                        output = output + f'-> Шрифт в тексте {head_str} не должен быть написан курсивом\n'
                        italic_checked = True
                        break

            
                if 'w:hAnsi' in str(par._p.xml):
                    if 'w:hAnsi="Times New Roman"' not in str(par._p.xml) and font_checked == False:
                        output = output + f'-> Тип шрифра в тексте {head_str} должен совпадать с типом шрифта в основном тексте\n'
                        font_checked = True
                elif 'w:hAnsiTheme' in str(par._p.xml):
                    if 'w:hAnsiTheme="Times New Roman"' not in str(par._p.xml) and font_checked == False:
                        output = output + f'-> Тип шрифра в тексте {head_str} должен совпадать с типом шрифта в основном тексте\n'
                        font_checked = True
                elif 'w:cs' in str(par._p.xml):
                    if 'w:cs="Times New Roman"' not in str(par._p.xml) and font_checked == False:
                        output = output + f'-> Тип шрифра в тексте {head_str} должен совпадать с типом шрифта в основном тексте\n'
                        font_checked = True



                for j in range(0, len(par.runs)):
                    if par.runs[j].font.underline == True and underline_checked == False:
                        output = output + f'-> Шрифт в тексте {head_str} не должен быть подчеркнут\n'
                        underline_checked = True
                        break

    if(output == "-- Проверка приложений --\n"):
        output = output + "-> OK\n"
    output = output + "\n"
    return output



            

            






