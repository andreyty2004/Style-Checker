import re
from bs4 import BeautifulSoup
import zipfile
from docx import Document


def tables(fpath = ""):
    doc = Document(fpath)
    output = "-- Проверка таблиц --\n"

    with zipfile.ZipFile(fpath, 'r') as zip_file:
        document_xml = zip_file.read('word/document.xml')

    soup = BeautifulSoup(document_xml, "xml")
    body = soup.find("w:document").find("w:body")

    global top_level_elements
    top_level_elements = body.find_all(["w:p", "w:tbl"], recursive = False)

    global page, isActive
    page = 1
    isActive = True # можно ли увеличивать page
    page_labels = []
    page_labels.append([0, 0])

    number_defined = False
    is_appendix = False

    total_table_count = 0 #будем искать соответствие между doc.paragraphs и top_level_elements

    global number, c
    number = None
    c = 0


    for i in range(0, len(top_level_elements)):
        top_element = top_level_elements[i]
        runs = top_element.find_all("w:r")

        if 'w:tbl' not in str(top_element):
            par = doc.paragraphs[i - total_table_count]

            if 'Heading 1' in par.style.name and len(par.runs) > 0:
                number_defined = False
                head_str = par.text

                if 'приложение' in head_str.lower():
                    is_appendix = True
                    start = head_str.lower().find('приложение') + len('приложение')
                    stroka = head_str[start:] #строка без 'приложение'
                    numbers = re.findall('([А-Я])', stroka)
                    if len(numbers) > 0:
                        number = numbers[0]
                        number_defined = True
                        c = 0
                else:
                    numbers = re.findall('(\d+)', head_str)
                    if len(numbers) > 0:
                        number = numbers[0]

                        if head_str.find(number) == 0:
                            number_defined = True
                            c = 0


        if 'w:tbl' in str(top_element):
            total_table_count += 1
            if number_defined:
                c += 1

            process_title(i, page, number_defined, number, c, is_appendix)
            if is_appendix == False:
                mentioned_above(page_labels, i, number_defined, number, c, page)


            boarders = False
            if len(top_element.find_all("w:tblBorders")) != 0:
                if "w:sz=\"8\"" in str(top_element.find_all("w:tblBorders")[0]):
                    boarders = True

            if boarders == False:
                output = output + f'-> Столбцы и строки таблицы на странице {page} ограничивают сплошными линиями толщиной 0,1 мм (1 pt)\n'

            check_font_size(top_element, page) #смотрим размер шрифта в таблице

            check_diagonals(top_element, page) #смотрим что нет диагоналей

        for j in range(0, len(runs)):
            page_increment(i, j, page_labels)

    if(output == "-- Проверка таблиц --\n"):
        output = output + "-> OK\n"
    output = output + "\n"
    return output

def page_increment(i, j, page_labels):
    global isActive
    global page

    runs_i = top_level_elements[i].find_all("w:r")
    if i != len(top_level_elements) - 1:
        runs_i_1 = top_level_elements[i + 1].find_all("w:r")

    if j != len(runs_i) - 1:
        if "w:br" in str(runs_i[j]) and "w:lastRenderedPageBreak" in str(runs_i[j + 1]):
            page += 1
            page_labels.append([i, j])
            isActive = False
            return 
        
    if i != len(top_level_elements) - 1 and len(runs_i_1) > 0 and j == len(runs_i) - 1:
        if "w:br" in str(runs_i[j]) and "w:lastRenderedPageBreak" in str(runs_i_1[0]):
            page += 1
            page_labels.append([i, j])
            isActive = False
            return

    if "w:br" in str(runs_i[j]):
        page += 1
        page_labels.append([i, j])
        return

    if "w:lastRenderedPageBreak" in str(runs_i[j]) and isActive:
        page += 1
        page_labels.append([i, j])
        return

    if "w:lastRenderedPageBreak" in str(runs_i[j]):
        isActive = True  
        return

def check_font_size(top_element, page):
    ftext = ""
    stroka = str(top_element)
    if (stroka.count('<w:sz w:val="28"/>') == 0 and 
        stroka.count('<w:sz w:val="24"/>') == 0 and 
        stroka.count('<w:sz w:val="20"/>') == 0):
            ftext = f'Размер шрифра в таблице на странице {page} может быть 10 pt, или 12 pt, или 14 pt\n'
    return ftext


def check_diagonals(top_element, page):
    ftext = ""
    cells = top_element.find_all("w:tc")
    for cell in cells:
        edge = cell.find("w:tcPr").find_all("w:tcBorders")
        if len(edge) != 0:
            edge = edge[0]
            tl2br = edge.find("w:tl2br")
            if "w:val=\"single\"" in str(tl2br):
                ftext = ftext + f"Ошибка в таблице на странице {page}. Разделять заголовки и подзаголовки в столбцах и строках таблицы диагональными линиями не допускается\n"
                return ftext
    return ftext


    
def check_numbering_in_text(index_of_mention, number, c, page):
    ftext = ""
    number_in_text = number + '.' + str(c)

    runs = top_level_elements[index_of_mention].find_all("w:r")
    temp = []

    for run in runs:
        w_ts = run.find_all("w:t")
        for w_t in w_ts:
            temp.append(w_t.text)

    text = ''.join(temp)

    if number_in_text not in text:
        ftext = f'-> Ошибка в номере в упоминании таблицы на странице {page} в тексте перед таблицей. Таблицы в разделах нумеруются по схеме «номер раздела – точка – номер таблицы»\n'
    return ftext


def mentioned_above(page_labels, i, number_defined, number, c, page):
    ftext = ""
    index_of_mention = -1

    start = -2
    if len(page_labels) == 1:
        start = -1

    temp = []

    i_0 = page_labels[start][0]
    j_0 = page_labels[start][1]

    runs_i_0 = top_level_elements[i_0].find_all("w:r")

    if "REF _Ref" in str(top_level_elements[i_0]):
        index_of_mention = i_0
    if "SEQ Таблица" not in str(top_level_elements[i_0]):
        for m in range(j_0, len(runs_i_0)):
            w_ts = runs_i_0[m].find_all("w:t")
            for w_t in w_ts:
                temp.append(w_t.text)

    for n in range(i_0 + 1, i):
        runs = top_level_elements[n].find_all("w:r")

        if "REF _Ref" in str(top_level_elements[n]):
            index_of_mention = n
        if "SEQ Таблица" not in str(top_level_elements[n]):
            for m in range(0, len(runs)):
                w_ts = runs[m].find_all("w:t")
                for w_t in w_ts:
                    temp.append(w_t.text)



    search_string = ''.join(temp)


    if 'таблиц' not in search_string.lower():
        ftext = f"Таблица на странице {page} должна располагаться непосредственно после текста, в котором она упоминается впервые, или на следующей странице. (не найдено упоминание в тексте)"
    else:
        if  index_of_mention == -1:
            ftext = ftext + f'Ссылка на таблицу на странице {page} в тексте должна являться перекрестной ссылкой (при нажатии на нее переносит на рисунок)'
        elif number_defined:
            ftext = ftext + check_numbering_in_text(index_of_mention, number, c, page)
    return ftext


def check_numbering(title_text, number, c, is_appendix, page):
    ftext = ""
    number_of_title = number + '.' + str(c)

    if number_of_title not in title_text:
        if is_appendix:
            ftext = f'-> Ошибка в подписи к таблице на странице {page}. Таблицы в приложениях нумеруются по схеме «номер приложения – точка – номер таблицы».\n'
        else:
            ftext = ftext + f'-> Ошибка в подписи к таблице на странице {page}. Таблицы в разделах нумеруются по схеме «номер раздела – точка – номер таблицы».\n'
    return ftext



def process_title(i, page, number_defined, number, c, is_appendix):
    ftext = ""
    for j in range(1, 4):
        par = top_level_elements[i - j]
        if "SEQ Таблица" in str(par):
            runs = par.find_all("w:r")
            title_list = []
            for run in runs:
                w_ts = run.find_all("w:t")
                for w_t in w_ts:
                    title_list.append(w_t.text)

            title_text = ''.join(title_list)

            if (title_text.lower().find("таблица") == 0 and "w:ind" not in str(par) and 
                "w:jc" not in str(par)) == False:
                ftext = f'-> Название таблицы на странице {page} должно помещаться на отдельной строке слева, без абзацного отступа\n'
                

                
            if 'таблица' in title_text.lower():
                if number_defined:
                    check_numbering(title_text, number, c, is_appendix, page)
            else:
                ftext = ftext + f'-> В подписи к таблице на странице {page} должно присутствовать слово "таблица"\n'
            return ftext
    ftext = ftext + f'-> Не найдена подпись к таблице на странице {page}. (выделить таблицу => пкм => вставить название)\n'
    return ftext

